from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "HUONG-DAN-HOAN-THIEN-DO-LUONG-HIEU-SUAT.md"
OUTPUT = ROOT / "docs" / "Huong-dan-hoan-thien-do-luong-hieu-suat-Dorae-Coffee.docx"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def clean_inline(text):
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"<((?:https?://|mailto:)[^>]+)>", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def add_table(document, rows):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index in range(width):
            value = clean_inline(values[column_index]) if column_index < len(values) else ""
            cells[column_index].text = value
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[column_index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if row_index == 0:
                set_cell_shading(cells[column_index], "5B3924")
            elif row_index % 2 == 0:
                set_cell_shading(cells[column_index], "F5EFE9")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_colors = {
        "Title": (75, 45, 29),
        "Heading 1": (75, 45, 29),
        "Heading 2": (142, 76, 42),
        "Heading 3": (142, 76, 42),
    }
    for style_name, color in heading_colors.items():
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = True

    styles["Title"].font.size = Pt(25)
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)

    header = section.header.paragraphs[0]
    header.text = "DORAE COFFEE  |  HƯỚNG DẪN ĐO LƯỜNG HIỆU SUẤT"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 100, 88)

    add_page_number(section.footer.paragraphs[0])


def build_document():
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    document = Document()
    configure_document(document)

    title = clean_inline(lines[0].removeprefix("# "))
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Website: https://doraecoffee.io.vn/")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 76, 48)
    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Tài liệu bàn giao triển khai, kiểm thử và vận hành").italic = True
    document.add_page_break()

    document.add_heading("Mục lục nội dung", level=1)
    for line in lines:
        if line.startswith("## "):
            document.add_paragraph(clean_inline(line[3:]), style="List Number")
    document.add_page_break()

    index = 1
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.6)
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(8)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                values = [part.strip() for part in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
                    table_lines.append(values)
                index += 1
            add_table(document, table_lines)
            continue

        if line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=2)
        elif line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=1)
        elif re.match(r"^\d+\. ", line):
            document.add_paragraph(clean_inline(re.sub(r"^\d+\. ", "", line)), style="List Number")
        elif line.startswith("- [ ] "):
            document.add_paragraph("☐ " + clean_inline(line[6:]), style="List Bullet")
        elif line.startswith("- "):
            document.add_paragraph(clean_inline(line[2:]), style="List Bullet")
        elif line.startswith("  - "):
            document.add_paragraph(clean_inline(line[4:]), style="List Bullet 2")
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            run = paragraph.add_run(clean_inline(line[2:]))
            run.italic = True
            run.font.color.rgb = RGBColor(92, 75, 65)
        elif line.strip():
            document.add_paragraph(clean_inline(line))
        index += 1

    properties = document.core_properties
    properties.title = title
    properties.subject = "Hoàn thiện hệ thống GA4, GTM, Search Console và dashboard KPI"
    properties.author = "Dorae Coffee"
    properties.keywords = "GA4, GTM, Search Console, KPI, SEO, UTM, Dorae Coffee"

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
